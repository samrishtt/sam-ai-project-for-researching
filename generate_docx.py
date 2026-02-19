
import os
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_html(raw_html):
    # Very basic html to plain text for docx (better to use a real parser but let's keep it simple)
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', raw_html)
    return cleantext

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Title
    if lines and lines[0].startswith('# '):
        title = lines[0][2:].strip()
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = lines[1:]

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # GLOBAL CLEANING: Remove math symbols from all lines before processing
        line = line.replace('$$', '').replace('$', '')
        line = line.replace('\\tau', 'tau').replace('\\lambda', 'lambda')
        line = line.replace('\\prod', 'Product').replace('\\cdot', 'x')
        line = line.replace('\\left', '').replace('\\right', '')
        line = line.replace('\\{', '{').replace('\\}', '}')
        line = line.replace('**', '').replace('*', '')

        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('!['):
            match = re.search(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('- ') or line.startswith('· '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line[re.search(r'\d+\.\s*', line).end():], style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)

if __name__ == "__main__":
    md_to_docx('research_paper.md', 'research_paper.docx')
    print("Successfully generated research_paper.docx")
